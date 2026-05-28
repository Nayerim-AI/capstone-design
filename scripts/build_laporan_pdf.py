#!/usr/bin/env python3
"""Build Google-Docs-style PDF from laporan_skripsi Markdown files.
Uses python-docx to create .docx, then LibreOffice to convert to PDF."""

import subprocess
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

DOCS = Path(__file__).resolve().parent.parent / "docs" / "laporan_skripsi"
OUT  = DOCS / "LAPORAN_SKRIPSI"


def setup_styles(doc):
    """Configure document for Google Docs / Buku Skripsi style."""
    # Page margins: 3cm top/bottom, 3cm left, 3cm right
    for section in doc.sections:
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)

    # Normal style
    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(12)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)

    # Heading styles
    for level, (size, bold, space_before, space_after) in {
        1: (Pt(14), True, Pt(24), Pt(12)),
        2: (Pt(12), True, Pt(18), Pt(6)),
        3: (Pt(12), True, Pt(12), Pt(6)),
    }.items():
        h = doc.styles[f'Heading {level}']
        h.font.name = 'Times New Roman'
        h.font.size = size
        h.font.bold = bold
        h.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        h.paragraph_format.space_before = space_before
        h.paragraph_format.space_after = space_after


def add_title_page(doc):
    """Add title page (Halaman Judul)."""
    for _ in range(6):
        doc.add_paragraph('')

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('DVB-T2 COVERAGE ANALYZER PORTABLE')
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Times New Roman'

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = subtitle.add_run('Analisis Cakupan Sinyal DVB-T2\nBerbasis RTL-SDR dan GPS\nKontrol via Telegram Bot')
    r2.font.size = Pt(12)
    r2.font.name = 'Times New Roman'

    for _ in range(4):
        doc.add_paragraph('')

    info_lines = [
        ('LAPORAN KULIAH KERJA NYATA (KKN)', False),
        ('', False),
        ('Nizar', False),
        ('12345678', False),
        ('', False),
        ('Program Studi Teknik Informatika', False),
        ('Fakultas Teknologi Informasi', False),
        ('Universitas XYZ', False),
        ('', False),
        ('2026', False),
    ]
    for text, bold in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'

    doc.add_page_break()


def parse_md_to_docx(md_path, doc):
    """Parse a simple Markdown file and add to docx."""
    content = md_path.read_text(encoding='utf-8')
    lines = content.split('\n')

    i = 0
    in_code_block = False
    code_buffer = []

    while i < len(lines):
        line = lines[i]

        # Code block start/end
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block — add code
                code_text = '\n'.join(code_buffer)
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        stripped = line.strip()

        # Empty line
        if not stripped:
            i += 1
            continue

        # Headings
        if stripped.startswith('### '):
            doc.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith('# '):
            doc.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith('---'):
            doc.add_paragraph('')
        else:
            # Regular paragraph — handle bold **text**
            p = doc.add_paragraph()
            parts = re.split(r'\*\*(.*?)\*\*', stripped)
            for j, part in enumerate(parts):
                if not part:
                    continue
                run = p.add_run(part)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                if j % 2 == 1:
                    run.bold = True

        i += 1


def main():
    OUT.mkdir(exist_ok=True, parents=True)

    doc = Document()
    setup_styles(doc)
    add_title_page(doc)

    files = [
        '00_Daftar_Isi.md',
        'CODING_PROGRAM.md',
        'LANGKAH_RUNNING_PROGRAM.md',
    ]

    for fname in files:
        src = DOCS / fname
        if not src.exists():
            print(f"[WARN] {src} not found, skipping.")
            continue
        print(f"→ Reading {fname} …")
        parse_md_to_docx(src, doc)

    docx_path = OUT / "LAPORAN_SKRIPSI.docx"
    doc.save(str(docx_path))
    print(f"  .docx → {docx_path}")

    # Convert to PDF with LibreOffice headless
    pdf_path = OUT / "LAPORAN_SKRIPSI.pdf"
    print(f"→ Converting to PDF …")
    result = subprocess.run([
        'libreoffice', '--headless', '--convert-to', 'pdf',
        '--outdir', str(OUT),
        str(docx_path),
    ], capture_output=True, text=True)
    if result.returncode != 0:
        print(f"  [ERROR] libreoffice: {result.stderr}")
    else:
        print(f"  .pdf  → {pdf_path}")

    # Merge all PDFs if multiple exist
    pdfs = sorted(OUT.glob("*.pdf"))
    if len(pdfs) > 1:
        print(f"→ Found {len(pdfs)} PDFs, will send individually.")
    print("\n✅ Done.")


if __name__ == '__main__':
    main()